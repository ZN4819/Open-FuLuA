from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Literal

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from ... import database
from ...config import settings
from ..docx_analyzer import analyze_docx
from ..template_profile import load_template_profile
from .fields import BookmarkWriter, add_complex_field
from .images import add_section_images, build_figure_refs
from .styles import apply_run_font, configure_section, set_paragraph_format
from .tables import add_assessment_table


ExportMode = Literal["editable", "final"]


class DocxGenerationError(RuntimeError):
    """DOCX 生成失败。"""


def generate_project_docx(project_id: int, mode: ExportMode = "editable") -> Path:
    if mode not in {"editable", "final"}:
        raise DocxGenerationError("导出模式必须是 editable 或 final。")

    project = database.get_project_by_id(project_id)
    if project is None:
        raise DocxGenerationError("项目不存在。")

    profile = load_template_profile()
    document = Document()
    _configure_document(document, profile)
    bookmark_writer = BookmarkWriter()

    sections = database.list_sections(project_id)
    for index, section in enumerate(sections):
        if index == 0:
            configure_section(document.sections[0], profile)
            _add_appendix_title(document, profile)
        else:
            configure_section(document.add_section(WD_SECTION.NEW_PAGE), profile)

        _add_section_title(document, section, profile)
        _add_table_caption(document, section, profile, bookmark_writer)

        images = database.list_evidence_images(project_id, section["code"])
        figure_refs = build_figure_refs(section["code"], images)
        rows = database.list_assessment_rows(section["id"])
        add_assessment_table(document, section, rows, profile, mode, figure_refs)
        add_section_images(document, section["code"], images, profile, bookmark_writer, figure_refs)

    export_path = _export_path(project_id, mode)
    document.save(export_path)
    _validate_generated_docx(export_path)
    return export_path


def _configure_document(document: Document, profile: dict[str, Any]) -> None:
    body = profile["typography"]["body"]
    normal = document.styles["Normal"]
    normal.font.name = body.get("ascii_font", "Times New Roman")
    normal.font.size = Pt(float(body["size_pt"]))
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), body.get("east_asia_font", "宋体"))

    settings_element = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings_element.append(update_fields)


def _add_appendix_title(document: Document, profile: dict[str, Any]) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph, profile, "appendix_title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("附录A测评结果记录")
    apply_run_font(run, profile, "appendix_title")


def _add_section_title(document: Document, section: Any, profile: dict[str, Any]) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph, profile, "section_title")
    run = paragraph.add_run(section["title"])
    apply_run_font(run, profile, "section_title")


def _add_table_caption(
    document: Document,
    section: Any,
    profile: dict[str, Any],
    bookmark_writer: BookmarkWriter,
) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph, profile, "caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table_number = section["code"].split("-")[-1]
    bookmark_name = f"tbl_{section['code'].replace('-', '_')}"
    bookmark_id = bookmark_writer.start(paragraph, bookmark_name)
    prefix = paragraph.add_run("表A-")
    apply_run_font(prefix, profile, "caption")
    add_complex_field(paragraph, "SEQ AppendixTable", table_number, lambda run: apply_run_font(run, profile, "caption"))
    bookmark_writer.end(paragraph, bookmark_id)

    table_title = section["table_title"]
    expected_prefix = f"表{section['code']}"
    suffix = table_title[len(expected_prefix):] if table_title.startswith(expected_prefix) else f" {table_title}"
    if suffix:
        suffix_run = paragraph.add_run(suffix)
        apply_run_font(suffix_run, profile, "caption")


def _export_path(project_id: int, mode: ExportMode) -> Path:
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    export_dir = settings.storage_path / "exports" / str(project_id)
    export_dir.mkdir(parents=True, exist_ok=True)
    return export_dir / f"appendix_a_project_{project_id}_{mode}_{timestamp}.docx"


def _validate_generated_docx(path: Path) -> None:
    analysis = analyze_docx(path)
    if analysis.sections != 8:
        raise DocxGenerationError(f"DOCX 分节数量不正确：{analysis.sections}")
    if analysis.tables != 8:
        raise DocxGenerationError(f"DOCX 表格数量不正确：{analysis.tables}")
    if analysis.missing_ref_targets:
        missing = ", ".join(analysis.missing_ref_targets)
        raise DocxGenerationError(f"DOCX 引用目标缺失：{missing}")
