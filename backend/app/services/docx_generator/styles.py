from __future__ import annotations

from typing import Any

from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.section import Section
from docx.shared import Cm, Inches, Pt
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def configure_section(section: Section, profile: dict[str, Any]) -> None:
    page = profile["page"]
    margins = page["margin_cm"]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(float(page["width_in"]))
    section.page_height = Inches(float(page["height_in"]))
    section.top_margin = Cm(float(margins["top"]))
    section.right_margin = Cm(float(margins["right"]))
    section.bottom_margin = Cm(float(margins["bottom"]))
    section.left_margin = Cm(float(margins["left"]))
    section.header_distance = Cm(float(margins["header"]))
    section.footer_distance = Cm(float(margins["footer"]))


def apply_run_font(run: Run, profile: dict[str, Any], role: str = "body", bold: bool | None = None) -> None:
    token = profile["typography"][role]
    font = run.font
    font.name = token.get("ascii_font", "Times New Roman")
    font.size = Pt(float(token["size_pt"]))
    set_run_complex_script_size(run, int(round(float(token["size_pt"]) * 2)))
    effective_bold = None
    if bold is not None:
        effective_bold = bool(bold)
    elif "bold" in token:
        effective_bold = bool(token["bold"])
    if effective_bold is not None:
        font.bold = effective_bold
        set_complex_script_bold(run, effective_bold)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), token.get("east_asia_font", "宋体"))
    if token.get("complex_script_font"):
        run._element.rPr.rFonts.set(qn("w:cs"), token["complex_script_font"])
    if token.get("caps"):
        set_run_caps(run, True)
    if "character_spacing_twips" in token:
        set_run_character_spacing(run, int(token["character_spacing_twips"]))
    if "kern_half_points" in token:
        set_run_kerning(run, int(token["kern_half_points"]))


def set_complex_script_bold(run: Run, enabled: bool) -> None:
    r_pr = run._element.get_or_add_rPr()
    bold_cs = r_pr.find(qn("w:bCs"))
    if bold_cs is None:
        bold_cs = OxmlElement("w:bCs")
        r_pr.append(bold_cs)
    if enabled:
        bold_cs.attrib.pop(qn("w:val"), None)
    else:
        bold_cs.set(qn("w:val"), "0")


def set_run_caps(run: Run, enabled: bool) -> None:
    r_pr = run._element.get_or_add_rPr()
    caps = r_pr.find(qn("w:caps"))
    if caps is None:
        caps = OxmlElement("w:caps")
        r_pr.append(caps)
    if enabled:
        caps.attrib.pop(qn("w:val"), None)
    else:
        caps.set(qn("w:val"), "0")


def set_run_character_spacing(run: Run, value_twips: int) -> None:
    r_pr = run._element.get_or_add_rPr()
    spacing = r_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        r_pr.append(spacing)
    spacing.set(qn("w:val"), str(value_twips))


def set_run_kerning(run: Run, value_half_points: int) -> None:
    r_pr = run._element.get_or_add_rPr()
    kern = r_pr.find(qn("w:kern"))
    if kern is None:
        kern = OxmlElement("w:kern")
        r_pr.append(kern)
    kern.set(qn("w:val"), str(value_half_points))


def set_run_complex_script_size(run: Run, value_half_points: int) -> None:
    r_pr = run._element.get_or_add_rPr()
    size = r_pr.find(qn("w:szCs"))
    if size is None:
        size = OxmlElement("w:szCs")
        r_pr.append(size)
    size.set(qn("w:val"), str(value_half_points))


def set_paragraph_format(paragraph: Paragraph, profile: dict[str, Any], role: str = "body") -> None:
    token = profile["typography"][role]
    alignment = token.get("alignment")
    if alignment == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == "both":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(float(token.get("spacing_before_twips", 0)) / 20)
    paragraph_format.space_after = Pt(float(token.get("spacing_after_twips", 0)) / 20)
    if token.get("line_twips"):
        paragraph_format.line_spacing = float(token["line_twips"]) / 240

    if "outline_level" in token:
        set_paragraph_outline_level(paragraph, int(token["outline_level"]))


def set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    profile: dict[str, Any],
    role: str = "body",
    bold: bool | None = None,
) -> None:
    paragraph.clear()
    set_paragraph_format(paragraph, profile, role)
    lines = (text or "").splitlines() or [""]
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        apply_run_font(run, profile, role, bold)


def set_cell_text(
    cell: _Cell,
    text: str,
    profile: dict[str, Any],
    role: str = "body",
    alignment: str = "left",
    bold: bool | None = None,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text, profile, role, bold)
    if alignment == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_cell_margins(cell: _Cell, margin_twips: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell: _Cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def configure_table_geometry(table: Table, column_widths_in: list[float], profile: dict[str, Any]) -> None:
    table.autofit = False
    total_width = sum(_inches_to_dxa(width) for width in column_widths_in)
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    colors = profile["colors"]
    set_table_borders(
        table,
        colors["border"],
        outer_size=int(colors.get("table_outer_border_size", 18)),
        inner_size=int(colors.get("table_inner_border_size", 4)),
    )

    for old_grid in tbl.findall(qn("w:tblGrid")):
        tbl.remove(old_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in column_widths_in:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(_inches_to_dxa(width)))
        tbl_grid.append(grid_col)
    tbl.insert(1, tbl_grid)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index < len(column_widths_in):
                width_dxa = _inches_to_dxa(column_widths_in[index])
                cell.width = Inches(column_widths_in[index])
                set_cell_width(cell, width_dxa)


def set_table_borders(table: Table, color: str, outer_size: int = 18, inner_size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(outer_size if side in {"top", "left", "bottom", "right"} else inner_size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_paragraph_outline_level(paragraph: Paragraph, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    outline = p_pr.first_child_found_in("w:outlineLvl")
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


def _inches_to_dxa(width: float) -> int:
    return int(round(float(width) * 1440))
