from __future__ import annotations

import math
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from ...config import settings
from .fields import BookmarkWriter, add_complex_field
from .styles import apply_run_font, set_paragraph_format


def build_figure_refs(section_code: str, images: list[Any]) -> dict[int, dict[str, str]]:
    refs: dict[int, dict[str, str]] = {}
    for index, image in enumerate(images, start=1):
        image_id = int(image["id"])
        refs[image_id] = {
            "bookmark": f"fig_{image_id}",
            "label": f"图{section_code}-{index}",
            "sequence": str(index),
        }
    return refs


def add_section_images(
    document: Document,
    section_code: str,
    images: list[Any],
    profile: dict[str, Any],
    bookmark_writer: BookmarkWriter,
    figure_refs: dict[int, dict[str, str]],
) -> None:
    if not images:
        return

    for image in images:
        image_id = int(image["id"])
        ref = figure_refs[image_id]
        image_path = settings.storage_path / image["file_path"]

        if not image_path.exists():
            paragraph = document.add_paragraph()
            set_paragraph_format(paragraph, profile, "caption")
            run = paragraph.add_run(f"[图片文件缺失：{image['original_name']}]")
            apply_run_font(run, profile, "caption")
            continue

        paragraph = document.add_paragraph()
        _format_figure_image_paragraph(paragraph)
        run = paragraph.add_run()
        width, height = _fit_image_size(image, ref, section_code, profile)
        inline_shape = run.add_picture(str(image_path), width=Inches(width), height=Inches(height))
        _set_image_alt_text(inline_shape, image["alt_text"] or image["caption"] or image["original_name"])

        caption = document.add_paragraph()
        set_paragraph_format(caption, profile, "caption")
        _format_figure_caption_paragraph(caption)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bookmark_id = bookmark_writer.start(caption, ref["bookmark"])
        prefix_run = caption.add_run(f"图{section_code}-")
        apply_run_font(prefix_run, profile, "caption")
        add_complex_field(caption, f"SEQ AppendixFigure_{section_code.replace('-', '_')}", ref["sequence"])
        bookmark_writer.end(caption, bookmark_id)

        caption_text = image["caption"].strip()
        if caption_text:
            suffix_run = caption.add_run(f" {caption_text}")
            apply_run_font(suffix_run, profile, "caption")


def _format_figure_image_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = paragraph.paragraph_format
    paragraph_format.page_break_before = True
    paragraph_format.keep_with_next = True
    paragraph_format.keep_together = True
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)


def _format_figure_caption_paragraph(paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.keep_together = True


def _fit_image_size(
    image: Any,
    ref: dict[str, str],
    section_code: str,
    profile: dict[str, Any],
) -> tuple[float, float]:
    usable_width = _usable_width(profile)
    usable_height = _usable_height(profile)
    max_width = min(float(profile["images"]["max_width_in"]), usable_width)
    caption_height = _caption_reserved_height(image, ref, section_code, profile, usable_width)
    max_height = max(1.0, usable_height - caption_height)
    aspect_ratio = _image_aspect_ratio(image)

    height_by_width = max_width / aspect_ratio
    if height_by_width <= max_height:
        return max_width, height_by_width
    return max_height * aspect_ratio, max_height


def _caption_reserved_height(
    image: Any,
    ref: dict[str, str],
    section_code: str,
    profile: dict[str, Any],
    usable_width: float,
) -> float:
    caption_profile = profile["typography"]["caption"]
    image_profile = profile["images"]
    font_size_pt = _positive_float(caption_profile.get("size_pt"), 12.0)
    line_height = _positive_float(caption_profile.get("line_twips"), font_size_pt * 24) / 1440
    space_before = _positive_float(caption_profile.get("spacing_before_twips"), 0) / 1440
    space_after = _positive_float(caption_profile.get("spacing_after_twips"), 0) / 1440
    caption_text = f"{ref.get('label', f'图{section_code}-')} {str(image['caption']).strip()}".strip()
    estimated_lines = _estimated_caption_lines(caption_text, usable_width, font_size_pt)
    estimated_height = space_before + space_after + line_height * estimated_lines
    return max(
        float(image_profile.get("caption_min_height_in", 0.45)),
        estimated_height + float(image_profile.get("caption_safety_height_in", 0.12)),
    )


def _estimated_caption_lines(text: str, usable_width: float, font_size_pt: float) -> int:
    character_width = max(font_size_pt / 72 * 0.9, 0.08)
    chars_per_line = max(int(usable_width / character_width), 1)
    return max(1, math.ceil(max(len(text), 1) / chars_per_line))


def _image_aspect_ratio(image: Any) -> float:
    pixel_width = _positive_float(image["pixel_width"], 0)
    pixel_height = _positive_float(image["pixel_height"], 0)
    if pixel_width > 0 and pixel_height > 0:
        return pixel_width / pixel_height

    display_width = _positive_float(image["display_width_in"], 0)
    display_height = _positive_float(image["display_height_in"], 0)
    if display_width > 0 and display_height > 0:
        return display_width / display_height

    return 1.0


def _usable_width(profile: dict[str, Any]) -> float:
    page = profile["page"]
    if page.get("usable_width_in"):
        return float(page["usable_width_in"])
    margins = page["margin_cm"]
    return float(page["width_in"]) - (float(margins["left"]) + float(margins["right"])) / 2.54


def _usable_height(profile: dict[str, Any]) -> float:
    page = profile["page"]
    margins = page["margin_cm"]
    return float(page["height_in"]) - (float(margins["top"]) + float(margins["bottom"])) / 2.54


def _positive_float(value: Any, fallback: float) -> float:
    if isinstance(value, (int, float)) and value > 0:
        return float(value)
    if isinstance(value, str):
        try:
            number = float(value)
        except ValueError:
            return fallback
        return number if number > 0 else fallback
    return fallback


def _set_image_alt_text(inline_shape, alt_text: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text)
