from __future__ import annotations

import re
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.table import _Cell, Table

from .content_controls import wrap_cell_paragraph_with_dropdown
from .fields import add_complex_field
from .styles import (
    apply_run_font,
    configure_table_geometry,
    set_cell_margins,
    set_cell_text,
    set_paragraph_format,
    shade_cell,
)


FIG_TOKEN_RE = re.compile(r"\[\[FIG:(\d+)\]\]")


def add_assessment_table(
    document: Document,
    section: Any,
    rows: list[Any],
    profile: dict[str, Any],
    mode: str,
    figure_refs: dict[int, dict[str, str]],
) -> Table:
    section_profile = _section_profile(profile, section["code"])
    table_profile = profile["tables"][section_profile["table_type"]]
    columns = table_profile["columns"]
    output_rows = _rows_with_calculated_unit_scores(rows or [_empty_row(section_profile["table_type"])], section_profile["table_type"])
    header_row_count = 2 if section_profile["table_type"] == "technical" else 1

    table = document.add_table(rows=header_row_count, cols=len(columns))
    table.style = None
    _add_header_rows(table, columns, section_profile["table_type"], profile)

    for row_index, row in enumerate(output_rows, start=1):
        table_row = table.add_row()
        for column_index, column in enumerate(columns):
            key = column["key"]
            cell = table_row.cells[column_index]
            _fill_body_cell(
                cell=cell,
                key=key,
                row=row,
                row_index=row_index,
                section_code=section["code"],
                profile=profile,
                mode=mode,
                figure_refs=figure_refs,
            )

    configure_table_geometry(table, [float(column["width_in"]) for column in columns], profile)
    _merge_repeated_unit_cells(table, output_rows, header_row_count, profile, columns)
    return table


def _add_header_rows(table: Table, columns: list[Any], table_type: str, profile: dict[str, Any]) -> None:
    if table_type == "technical":
        _add_technical_header_rows(table, columns, profile)
        return
    _add_management_header_row(table, columns, profile)


def _add_management_header_row(table: Table, columns: list[Any], profile: dict[str, Any]) -> None:
    header = table.rows[0]
    _mark_repeat_header(header)
    for index, column in enumerate(columns):
        if column["key"] == "unit_score":
            _set_score_header_cell(header.cells[index], column["label"], "management", profile)
        else:
            _set_header_cell(header.cells[index], column["label"], profile)


def _add_technical_header_rows(table: Table, columns: list[Any], profile: dict[str, Any]) -> None:
    for row in table.rows[:2]:
        _mark_repeat_header(row)
        for cell in row.cells:
            shade_cell(cell, profile["colors"]["table_header_fill"])
            set_cell_margins(cell)

    _set_header_cell(table.cell(0, 0).merge(table.cell(1, 0)), columns[0]["label"], profile)
    _set_header_cell(table.cell(0, 1).merge(table.cell(1, 1)), columns[1]["label"], profile)
    _set_header_cell(table.cell(0, 2).merge(table.cell(1, 2)), columns[2]["label"], profile)
    _set_header_cell(table.cell(0, 3).merge(table.cell(0, 6)), "量化指标", profile)
    _set_score_header_cell(table.cell(0, 7).merge(table.cell(1, 7)), columns[7]["label"], "technical", profile)

    for index in range(3, 7):
        _set_header_cell(table.cell(1, index), columns[index]["label"].replace(" ", ""), profile)


def _set_header_cell(cell: _Cell, text: str, profile: dict[str, Any]) -> None:
    set_cell_text(cell, text, profile, "table_header", "center")
    shade_cell(cell, profile["colors"]["table_header_fill"])


def _set_score_header_cell(cell: _Cell, text: str, table_type: str, profile: dict[str, Any]) -> None:
    _set_header_cell(cell, text, profile)
    paragraph = cell.add_paragraph()
    set_paragraph_format(paragraph, profile, "table_header")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph._p.append(_score_formula_xml(table_type))


def _fill_body_cell(
    cell: _Cell,
    key: str,
    row: Any,
    row_index: int,
    section_code: str,
    profile: dict[str, Any],
    mode: str,
    figure_refs: dict[int, dict[str, str]],
) -> None:
    if key == "record_text":
        _set_record_cell(cell, _value(row, key), profile, figure_refs)
        return

    if key == "unit":
        _set_unit_cell(cell, _value(row, key), profile)
        return

    if key in {"d", "a", "k"}:
        value = _value(row, key) or profile["content_controls"]["technical_metric"]["default"]
        _set_metric_cell(
            cell=cell,
            value=value,
            tag=f"{section_code.replace('-', '')}.row{row_index}.{key.upper()}",
            options=profile["content_controls"]["technical_metric"]["options"],
            profile=profile,
            mode=mode,
        )
        return

    if key == "compliance":
        value = _value(row, key) or profile["content_controls"]["management_compliance"]["default"]
        _set_metric_cell(
            cell=cell,
            value=value,
            tag=f"{section_code.replace('-', '')}.row{row_index}.compliance",
            options=profile["content_controls"]["management_compliance"]["options"],
            profile=profile,
            mode=mode,
        )
        return

    alignment = "center" if key in {"object_name", "object_score", "unit_score"} else "left"
    set_cell_text(cell, _value(row, key), profile, "body", alignment)


def _set_unit_cell(cell: _Cell, text: str, profile: dict[str, Any]) -> None:
    set_cell_text(cell, text, profile, "body", "center", bold=True)
    shade_cell(cell, profile["colors"].get("table_unit_fill", profile["colors"]["table_header_fill"]))


def _set_metric_cell(
    cell: _Cell,
    value: str,
    tag: str,
    options: list[str],
    profile: dict[str, Any],
    mode: str,
) -> None:
    set_cell_text(cell, value, profile, "body", "center")
    if mode == "editable":
        wrap_cell_paragraph_with_dropdown(cell, tag, value, options)


def _set_record_cell(
    cell: _Cell,
    text: str,
    profile: dict[str, Any],
    figure_refs: dict[int, dict[str, str]],
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_format(paragraph, profile, "body")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)

    cursor = 0
    for match in FIG_TOKEN_RE.finditer(text or ""):
        if match.start() > cursor:
            _append_text(paragraph, text[cursor:match.start()], profile)
        image_id = int(match.group(1))
        ref = figure_refs.get(image_id)
        if ref:
            add_complex_field(paragraph, f"REF {ref['bookmark']} \\h", ref["label"])
        else:
            _append_text(paragraph, match.group(0), profile)
        cursor = match.end()

    if cursor < len(text or ""):
        _append_text(paragraph, (text or "")[cursor:], profile)


def _append_text(paragraph, text: str, profile: dict[str, Any]) -> None:
    parts = text.splitlines()
    for index, part in enumerate(parts):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        apply_run_font(run, profile, "body")


def _score_formula_xml(table_type: str):
    if table_type == "management":
        formula = _subscript("S", "i,j")
    else:
        formula = (
            f"{_subscript('S', 'i,j')}"
            f"{_math_run('=')}"
            "<m:f>"
            "<m:num>"
            "<m:nary>"
            '<m:naryPr><m:chr m:val="∑"/><m:limLoc m:val="undOvr"/><m:supHide m:val="1"/></m:naryPr>'
            f"<m:sub>{_math_run('1≤k≤')}{_subscript('n', 'i,j')}</m:sub>"
            f"<m:e>{_subscript('S', 'i,j,k')}</m:e>"
            "</m:nary>"
            "</m:num>"
            f"<m:den>{_subscript('n', 'i,j')}</m:den>"
            "</m:f>"
        )
    return parse_xml(f'<m:oMathPara {nsdecls("m", "w")}><m:oMath>{formula}</m:oMath></m:oMathPara>')


def _subscript(base: str, subscript: str) -> str:
    return f"<m:sSub><m:e>{_math_run(base)}</m:e><m:sub>{_math_run(subscript)}</m:sub></m:sSub>"


def _math_run(text: str) -> str:
    return (
        "<m:r>"
        '<m:rPr><m:sty m:val="bi"/></m:rPr>'
        "<w:rPr>"
        '<w:rFonts w:ascii="Cambria Math" w:eastAsia="宋体" w:hAnsi="Cambria Math" w:cs="Times New Roman"/>'
        '<w:sz w:val="21"/><w:szCs w:val="21"/>'
        "</w:rPr>"
        f"<m:t>{text}</m:t>"
        "</m:r>"
    )


def _section_profile(profile: dict[str, Any], code: str) -> dict[str, Any]:
    for section in profile["sections"]:
        if section["code"] == code:
            return section
    raise ValueError(f"模板 profile 缺少章节：{code}")


def _empty_row(table_type: str) -> dict[str, str]:
    if table_type == "technical":
        return {"unit": "", "object_name": "", "record_text": "", "d": "/", "a": "/", "k": ""}
    return {"unit": "", "object_name": "", "record_text": "", "compliance": "不适用"}


def _value(row: Any, key: str) -> str:
    if isinstance(row, dict):
        value = row.get(key)
    elif hasattr(row, "keys") and key in row.keys():
        value = row[key]
    else:
        value = None
    return "" if value is None else str(value)


def _mark_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _rows_with_calculated_unit_scores(rows: list[Any], table_type: str) -> list[Any]:
    if table_type != "technical":
        return rows

    rows_by_unit: dict[str, list[Any]] = {}
    for row in rows:
        rows_by_unit.setdefault(_value(row, "unit").strip(), []).append(row)
    score_by_unit = {
        unit: _calculate_unit_score(unit_rows)
        for unit, unit_rows in rows_by_unit.items()
    }

    output_rows: list[dict[str, Any]] = []
    for row in rows:
        row_data = _row_to_dict(row)
        row_data["unit_score"] = score_by_unit.get(_value(row, "unit").strip(), "")
        output_rows.append(row_data)
    return output_rows


def _row_to_dict(row: Any) -> dict[str, Any]:
    if isinstance(row, dict):
        return dict(row)
    if hasattr(row, "keys"):
        return {key: row[key] for key in row.keys()}
    return {}


def _calculate_unit_score(rows: list[Any]) -> str:
    numeric_scores: list[Decimal] = []
    filled_scores = 0
    excluded_scores = 0
    for row in rows:
        score = _value(row, "object_score").strip()
        if not score:
            continue
        filled_scores += 1
        if score == "/":
            excluded_scores += 1
            continue
        try:
            numeric_scores.append(Decimal(score))
        except InvalidOperation:
            continue

    if numeric_scores:
        average = sum(numeric_scores) / Decimal(len(numeric_scores))
        return str(average.quantize(Decimal("0.0001"), rounding=ROUND_HALF_UP))
    if rows and filled_scores == len(rows) and excluded_scores == len(rows):
        return "/"
    return ""


def _merge_repeated_unit_cells(
    table: Table,
    rows: list[Any],
    header_row_count: int,
    profile: dict[str, Any],
    columns: list[Any],
) -> None:
    start_index = header_row_count
    previous_unit = None
    unit_score_column = next((index for index, column in enumerate(columns) if column["key"] == "unit_score"), None)
    for offset, row in enumerate(rows + [None], start=header_row_count):
        unit = _value(row, "unit") if row is not None else None
        if offset == header_row_count:
            previous_unit = unit
            start_index = header_row_count
            continue
        if unit != previous_unit:
            end_index = offset - 1
            if previous_unit and end_index > start_index:
                merged_cell = table.cell(start_index, 0).merge(table.cell(end_index, 0))
                _set_unit_cell(merged_cell, previous_unit, profile)
                if unit_score_column is not None:
                    score_cell = table.cell(start_index, unit_score_column).merge(table.cell(end_index, unit_score_column))
                    set_cell_text(score_cell, _value(rows[start_index - header_row_count], "unit_score"), profile, "body", "center")
            start_index = offset
            previous_unit = unit
