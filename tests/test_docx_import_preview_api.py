import json
import os
import sys
import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import HTTPException, UploadFile
from PIL import Image
from starlette.datastructures import Headers


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "backend"))

from app import database  # noqa: E402
from app.api.imports import get_docx_import as api_get_docx_import  # noqa: E402
from app.api.imports import upload_docx_import as api_upload_docx_import  # noqa: E402
from app.config import settings  # noqa: E402
from app.services.docx_generator import generate_project_docx  # noqa: E402


class DocxImportPreviewApiTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)
        self.root = Path(self.temp_dir.name)
        self.original_storage_path = settings.storage_path
        object.__setattr__(settings, "storage_path", self.root / "storage")
        os.environ["FULUA_DATABASE_PATH"] = str(self.root / "test.db")
        database.init_db()

    def tearDown(self) -> None:
        os.environ.pop("FULUA_DATABASE_PATH", None)
        object.__setattr__(settings, "storage_path", self.original_storage_path)

    def test_upload_generated_docx_returns_preview_job_and_parsed_json(self) -> None:
        project = self._make_project_with_content()
        exported = generate_project_docx(project["id"], "editable")
        upload = self._upload("XX系统附录A.docx", exported.read_bytes())

        job = api_upload_docx_import(upload)
        reloaded = api_get_docx_import(job.id)

        self.assertEqual(job.status, "preview_ready")
        self.assertTrue(job.can_create_project)
        self.assertEqual(job.suggested_project_name, "XX系统附录A")
        self.assertEqual(job.source_docx_path, f"imports/{job.id}/source.docx")
        self.assertEqual(job.parsed_json_path, f"imports/{job.id}/parsed.json")
        self.assertFalse(Path(job.source_docx_path).is_absolute())
        self.assertFalse(Path(job.parsed_json_path or "").is_absolute())
        self.assertEqual(job.summary["assessment_rows"], 1)
        self.assertEqual(job.summary["images"], 1)
        self.assertEqual(job.summary["references"], 1)
        self.assertEqual(job.sections[0].code, "A-1")
        self.assertEqual(job.sections[0].row_count, 1)
        self.assertEqual(job.sections[0].image_count, 1)
        self.assertEqual(job.sections[0].reference_count, 1)
        self.assertEqual(reloaded.id, job.id)
        self.assertEqual(reloaded.sections[0].image_count, 1)

        parsed_path = settings.storage_path / str(job.parsed_json_path)
        parsed_payload = json.loads(parsed_path.read_text(encoding="utf-8"))
        self.assertEqual(parsed_payload["sections"][0]["rows"][0]["record_text"], "查看登录策略，见 [[FIG:import:A-1-1]]。")
        self.assertEqual(parsed_payload["sections"][0]["images"][0]["figure_label"], "图A-1-1")

    def test_non_docx_upload_is_rejected_without_creating_job(self) -> None:
        upload = self._upload("not-docx.txt", b"not a docx", "text/plain")

        with self.assertRaises(HTTPException) as context:
            api_upload_docx_import(upload)

        self.assertEqual(context.exception.status_code, 400)
        self.assertEqual(database.get_docx_import_job(1), None)

    def test_corrupted_docx_returns_failed_job(self) -> None:
        upload = self._upload("broken.docx", b"not a zip package")

        job = api_upload_docx_import(upload)
        reloaded = api_get_docx_import(job.id)

        self.assertEqual(job.status, "failed")
        self.assertFalse(job.can_create_project)
        self.assertEqual(job.summary["errors"], 1)
        self.assertEqual(job.issues[0].code, "IMPORT_PARSE_FAILED")
        self.assertIn("损坏", job.error_message or "")
        self.assertEqual(reloaded.status, "failed")

    def test_preview_with_structure_errors_disables_project_creation(self) -> None:
        path = self.root / "partial.docx"
        document = Document()
        document.add_paragraph("附录A测评结果记录")
        document.add_paragraph("A-1 物理和环境安全")
        document.add_paragraph("表A-1物理和环境安全测评结果记录")
        table = document.add_table(rows=2, cols=8)
        for index, label in enumerate(["测评单元", "测评对象", "结果记录", "D", "A", "K", "测评对象评分", "测评单元得分"]):
            table.cell(0, index).text = label
        table.cell(1, 0).text = "身份鉴别"
        table.cell(1, 1).text = "机房"
        table.cell(1, 2).text = "现场查看。"
        table.cell(1, 3).text = "√"
        table.cell(1, 4).text = "√"
        table.cell(1, 5).text = "/"
        path.parent.mkdir(parents=True, exist_ok=True)
        document.save(path)

        job = api_upload_docx_import(self._upload("partial.docx", path.read_bytes()))
        issue_codes = {issue.code for issue in job.issues}

        self.assertEqual(job.status, "preview_ready")
        self.assertFalse(job.can_create_project)
        self.assertGreater(job.summary["errors"], 0)
        self.assertIn("IMPORT_MISSING_TABLE", issue_codes)
        self.assertEqual(job.sections[0].row_count, 1)

    def _make_project_with_content(self):
        project = database.create_project("DOCX 导入预览 API 测试")
        image = self._create_evidence_image(project["id"], "A-1")
        database.replace_section_rows(
            project_id=project["id"],
            code="A-1",
            rows=[
                {
                    "unit": "身份鉴别",
                    "object_name": "服务器",
                    "record_text": f"查看登录策略，见 [[FIG:{image['id']}]]。",
                    "metric_result": {
                        "d": "√",
                        "a": "√",
                        "k": "/",
                        "object_score": "1.0000",
                        "unit_score": "1.0000",
                    },
                    "cross_references": [
                        {
                            "target_image_id": image["id"],
                            "token": f"[[FIG:{image['id']}]]",
                            "display_text": "图A-1-1",
                        }
                    ],
                }
            ],
        )
        return project

    def _create_evidence_image(self, project_id: int, section_code: str):
        relative_path = Path("uploads") / str(project_id) / section_code / "preview.png"
        absolute_path = settings.storage_path / relative_path
        absolute_path.parent.mkdir(parents=True, exist_ok=True)
        Image.new("RGB", (600, 300), color=(255, 255, 255)).save(absolute_path, dpi=(150, 150))
        return database.create_evidence_image(
            project_id,
            section_code,
            {
                "file_path": relative_path.as_posix(),
                "original_name": "preview.png",
                "caption": "登录策略截图",
                "alt_text": "登录策略截图",
                "pixel_width": 600,
                "pixel_height": 300,
                "dpi_x": 150,
                "dpi_y": 150,
                "display_width_in": 4,
                "display_height_in": 2,
            },
        )

    @staticmethod
    def _upload(filename: str, content: bytes, content_type: str = "application/vnd.openxmlformats-officedocument.wordprocessingml.document") -> UploadFile:
        return UploadFile(
            file=BytesIO(content),
            filename=filename,
            headers=Headers({"content-type": content_type}),
        )


if __name__ == "__main__":
    unittest.main()