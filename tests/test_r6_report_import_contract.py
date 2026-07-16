from __future__ import annotations

import hashlib
import io
import os
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from fastapi import HTTPException, UploadFile
from starlette.datastructures import Headers

from app import database
from app.api.report_imports import (
    confirm_report_import,
    get_report_import,
    upload_report_import,
)
from app.main import app
from app.report_import import parser as report_import_parser
from app.report_import.schemas import ReportImportConfirmWrite
from app.services.docx_importer.tables import (
    FULL_REPORT_APPENDIX_TABLE_INDICES,
    parse_full_report_appendix_tables,
)


ROOT = Path(__file__).resolve().parents[1]
RUNTIME_TEMPLATE = ROOT / "templates" / "report" / "2023-2025.12.08" / "runtime_template.docx"


def _find_local_sample(environment_name: str) -> Path | None:
    configured = os.getenv(environment_name)
    if not configured:
        return None
    candidate = Path(configured)
    return candidate.resolve() if candidate.is_file() else None


LOCAL_CUSTOMER_REPORT = _find_local_sample("FULUA_R6_CUSTOMER_DOCX")


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        while chunk := stream.read(1024 * 1024):
            digest.update(chunk)
    return digest.hexdigest()


def _upload(filename: str, data: bytes) -> UploadFile:
    return UploadFile(
        filename=filename,
        file=io.BytesIO(data),
        headers=Headers(
            {
                "content-type": (
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            }
        ),
    )


def _rewrite_docx(source: bytes, replacements: dict[str, bytes]) -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(source)) as package, zipfile.ZipFile(
        output,
        "w",
        compression=zipfile.ZIP_DEFLATED,
    ) as rewritten:
        existing = {item.filename for item in package.infolist()}
        for item in package.infolist():
            if item.filename in replacements:
                rewritten.writestr(item, replacements[item.filename])
            else:
                rewritten.writestr(item, package.read(item))
        for name, data in replacements.items():
            if name not in existing:
                rewritten.writestr(name, data)
    return output.getvalue()


class R6ReportImportContractTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory()
        self.root = Path(self.temporary.name)
        self.previous = {
            name: os.environ.get(name)
            for name in ("FULUA_DATA_DIR", "FULUA_DATABASE_PATH", "FULUA_STORAGE_PATH")
        }
        os.environ["FULUA_DATA_DIR"] = str(self.root)
        os.environ["FULUA_DATABASE_PATH"] = str(self.root / "data" / "app.db")
        os.environ["FULUA_STORAGE_PATH"] = str(self.root / "storage")
        database.init_db()

    def tearDown(self) -> None:
        for name, value in self.previous.items():
            if value is None:
                os.environ.pop(name, None)
            else:
                os.environ[name] = value
        self.temporary.cleanup()

    def test_openapi_exposes_migration_only_workflow_and_appendix_copy(self) -> None:
        schema = app.openapi()
        expected_methods = {
            "/api/report-imports/docx": "post",
            "/api/report-imports/{job_id}": "get",
            "/api/report-imports/{job_id}/resolutions": "put",
            "/api/report-imports/{job_id}/confirm": "post",
            "/api/projects/{target_uuid}/report/appendix-a/copy": "post",
        }
        for path, method in expected_methods.items():
            with self.subTest(path=path, method=method):
                self.assertIn(path, schema["paths"])
                self.assertIn(method, schema["paths"][path])

        upload_parameters = schema["paths"]["/api/report-imports/docx"]["post"]["parameters"]
        mode = next(parameter for parameter in upload_parameters if parameter["name"] == "mode")
        self.assertEqual(mode["in"], "query")
        self.assertEqual(mode["schema"].get("default"), "migration")

    def test_roundtrip_is_explicitly_not_implemented_and_does_not_create_job(self) -> None:
        before = self._job_count()
        with self.assertRaises(HTTPException) as captured:
            upload_report_import(_upload("report.docx", RUNTIME_TEMPLATE.read_bytes()), "roundtrip")

        self.assertEqual(captured.exception.status_code, 501)
        detail = captured.exception.detail
        self.assertIsInstance(detail, dict)
        self.assertEqual(detail.get("code"), "ROUNDTRIP_NOT_IMPLEMENTED")
        self.assertEqual(self._job_count(), before)

    def test_non_docx_and_empty_upload_are_rejected_without_job(self) -> None:
        for filename, content in (("report.docm", b"x"), ("report.docx", b"")):
            with self.subTest(filename=filename, empty=not content):
                before = self._job_count()
                with self.assertRaises(HTTPException) as captured:
                    upload_report_import(_upload(filename, content), "migration")
                self.assertEqual(captured.exception.status_code, 400)
                self.assertEqual(self._job_count(), before)

    def test_path_traversal_package_is_audited_as_failed_and_cannot_confirm(self) -> None:
        output = io.BytesIO()
        with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as package:
            package.writestr("[Content_Types].xml", b"<Types/>")
            package.writestr("_rels/.rels", b"<Relationships/>")
            package.writestr("word/document.xml", b"<document/>")
            package.writestr("../outside.txt", b"must-not-escape")

        job = upload_report_import(_upload("unsafe.docx", output.getvalue()), "migration")

        self._assert_failed_job(job)
        self.assertFalse((self.root / "outside.txt").exists())
        with self.assertRaises(HTTPException) as captured:
            confirm_report_import(
                job.id,
                ReportImportConfirmWrite(
                    job_revision=job.job_revision,
                    project_name="禁止创建",
                    appendix_a_source="document",
                ),
            )
        self.assertEqual(captured.exception.status_code, 409)
        self.assertEqual(database.list_projects(), [])

    def test_activex_and_dde_packages_are_audited_as_failed(self) -> None:
        original = RUNTIME_TEMPLATE.read_bytes()
        with zipfile.ZipFile(io.BytesIO(original)) as package:
            document_xml = package.read("word/document.xml")
        dde_xml = document_xml.replace(
            b"</w:body>",
            b'<w:p><w:r><w:instrText>DDEAUTO "cmd" "/c calc"</w:instrText></w:r></w:p></w:body>',
            1,
        )
        split_dde_xml = document_xml.replace(
            b"</w:body>",
            (
                b'<w:p><w:r><w:instrText>D</w:instrText></w:r>'
                b'<w:r><w:instrText>DE "cmd" "/c calc"</w:instrText></w:r></w:p>'
                b"</w:body>"
            ),
            1,
        )
        split_ddeauto_xml = document_xml.replace(
            b"</w:body>",
            (
                b'<w:p><w:r><w:fldChar w:fldCharType="begin"/></w:r>'
                b'<w:r><w:instrText>DDE</w:instrText></w:r>'
                b'<w:r><w:instrText>AUTO "cmd" "/c calc"</w:instrText></w:r>'
                b'<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
                b'<w:r><w:t>result</w:t></w:r>'
                b'<w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>'
                b"</w:body>"
            ),
            1,
        )
        variants = {
            "activex.docx": _rewrite_docx(original, {"word/activeX/activeX1.bin": b"x"}),
            "dde.docx": _rewrite_docx(original, {"word/document.xml": dde_xml}),
            "split-dde.docx": _rewrite_docx(original, {"word/document.xml": split_dde_xml}),
            "split-ddeauto.docx": _rewrite_docx(
                original,
                {"word/document.xml": split_ddeauto_xml},
            ),
        }
        for filename, payload in variants.items():
            with self.subTest(filename=filename):
                job = upload_report_import(_upload(filename, payload), "migration")
                self._assert_failed_job(job)

    def test_unapproved_embedding_and_attached_template_are_audited_as_failed(self) -> None:
        original = RUNTIME_TEMPLATE.read_bytes()
        with zipfile.ZipFile(io.BytesIO(original)) as package:
            document_relationships = package.read("word/_rels/document.xml.rels")
        external_relationships = document_relationships.replace(
            b"</Relationships>",
            (
                b'<Relationship Id="rIdUnsafeExternal" '
                b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" '
                b'Target="https://example.invalid/unsafe" TargetMode="External"/>'
                b"</Relationships>"
            ),
            1,
        )
        attached_relationship = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rIdUnsafe" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/attachedTemplate" Target="https://example.invalid/template.dotm" TargetMode="External"/>
</Relationships>"""
        variants = {
            "embedding.docx": _rewrite_docx(
                original,
                {"word/embeddings/unknown.bin": b"unapproved embedding"},
            ),
            "attached-template.docx": _rewrite_docx(
                original,
                {"word/_rels/settings.xml.rels": attached_relationship},
            ),
            "external-link.docx": _rewrite_docx(
                original,
                {"word/_rels/document.xml.rels": external_relationships},
            ),
        }
        for filename, payload in variants.items():
            with self.subTest(filename=filename):
                self._assert_failed_job(
                    upload_report_import(_upload(filename, payload), "migration")
                )

    def test_only_exactly_allowlisted_embedding_is_ignored_with_warning(self) -> None:
        part_name = "word/embeddings/oleObject1.bin"
        fixture = b"synthetic known embedding"
        payload = _rewrite_docx(
            RUNTIME_TEMPLATE.read_bytes(),
            {part_name: fixture},
        )
        digest = hashlib.sha256(fixture).hexdigest()

        with patch.dict(
            report_import_parser.KNOWN_IGNORED_OLE,
            {part_name: digest},
            clear=True,
        ):
            job = upload_report_import(_upload("known-embedding.docx", payload), "migration")

        self.assertNotEqual(job.status, "failed")
        self.assertTrue(any(issue.code == "KNOWN_OLE_IGNORED" for issue in job.issues))

    def test_external_hyperlink_owned_by_comments_is_ignored_with_warning(self) -> None:
        original = RUNTIME_TEMPLATE.read_bytes()
        with zipfile.ZipFile(io.BytesIO(original)) as package:
            content_types = package.read("[Content_Types].xml")
        content_types = content_types.replace(
            b"</Types>",
            (
                b'<Override PartName="/word/comments.xml" '
                b'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
                b"</Types>"
            ),
            1,
        )
        comments = (
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        comment_relationships = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" Target="https://example.invalid/reference" TargetMode="External"/>
</Relationships>"""
        payload = _rewrite_docx(
            original,
            {
                "[Content_Types].xml": content_types,
                "word/comments.xml": comments,
                "word/_rels/comments.xml.rels": comment_relationships,
            },
        )

        job = upload_report_import(_upload("comment-link.docx", payload), "migration")

        self.assertNotEqual(job.status, "failed")
        self.assertTrue(any(issue.severity == "warning" for issue in job.issues))

    def test_unknown_structure_is_persisted_but_never_confirmable(self) -> None:
        path = self.root / "unknown.docx"
        document = Document()
        document.add_heading("未知报告", 1)
        document.add_table(rows=2, cols=2)
        document.save(path)

        job = upload_report_import(_upload(path.name, path.read_bytes()), "migration")

        self._assert_failed_job(job)
        self.assertIsNone(job.detected_edition)
        self.assertIsNone(job.detected_revision)

    def test_runtime_template_matches_frozen_fingerprint_without_source_mutation(self) -> None:
        before = (_sha256(RUNTIME_TEMPLATE), RUNTIME_TEMPLATE.stat().st_mtime_ns)

        job = upload_report_import(
            _upload(RUNTIME_TEMPLATE.name, RUNTIME_TEMPLATE.read_bytes()),
            "migration",
        )
        reloaded = get_report_import(job.id)

        self.assertNotEqual(job.status, "failed")
        self.assertEqual(job.detected_edition, "2023")
        self.assertEqual(job.detected_revision, "2025-12-08")
        self.assertTrue(job.fingerprint.matched)
        self.assertEqual(reloaded.id, job.id)
        with database.connect() as connection:
            stored = connection.execute(
                "SELECT source_docx_path, source_sha256 FROM report_import_jobs WHERE id = ?",
                (job.id,),
            ).fetchone()
        stored_relative_path = Path(stored["source_docx_path"])
        self.assertFalse(stored_relative_path.is_absolute())
        self.assertEqual(stored["source_sha256"], before[0])
        self.assertEqual(_sha256(self.root / "storage" / stored_relative_path), before[0])
        self.assertEqual((_sha256(RUNTIME_TEMPLATE), RUNTIME_TEMPLATE.stat().st_mtime_ns), before)

    def test_full_report_appendix_uses_frozen_table_identities(self) -> None:
        parsed = parse_full_report_appendix_tables(RUNTIME_TEMPLATE)

        self.assertEqual([section.code for section in parsed.sections], list(FULL_REPORT_APPENDIX_TABLE_INDICES))
        self.assertEqual(parsed.summary["parsed_sections"], 8)
        self.assertFalse(any(issue.severity == "error" for issue in parsed.issues))
        for section in parsed.sections:
            self.assertTrue(section.rows)
            self.assertEqual(
                {row.source_table_index for row in section.rows},
                {FULL_REPORT_APPENDIX_TABLE_INDICES[section.code]},
            )

    @unittest.skipUnless(
        LOCAL_CUSTOMER_REPORT is not None,
        "仅在本地客户复核版存在或设置 FULUA_R6_CUSTOMER_DOCX 时运行",
    )
    def test_local_customer_report_is_read_only_and_matches_supported_family(self) -> None:
        assert LOCAL_CUSTOMER_REPORT is not None
        before = (
            _sha256(LOCAL_CUSTOMER_REPORT),
            LOCAL_CUSTOMER_REPORT.stat().st_size,
            LOCAL_CUSTOMER_REPORT.stat().st_mtime_ns,
        )

        job = upload_report_import(
            _upload("local-acceptance-source.docx", LOCAL_CUSTOMER_REPORT.read_bytes()),
            "migration",
        )

        self.assertNotEqual(job.status, "failed")
        self.assertEqual(job.detected_edition, "2023")
        self.assertEqual(job.detected_revision, "2025-12-08")
        self.assertTrue(job.fingerprint.matched)
        self.assertFalse(Path(job.original_name).is_absolute())
        self.assertFalse(job.summary["document_appendix"]["available"])
        with self.assertRaises(HTTPException) as captured:
            confirm_report_import(
                job.id,
                ReportImportConfirmWrite(
                    job_revision=job.job_revision,
                    project_name="本地只读验收",
                    appendix_a_source="document",
                ),
            )
        self.assertEqual(captured.exception.detail["code"], "APPENDIX_A_INCOMPLETE")
        self.assertEqual(database.list_projects(), [])
        self.assertEqual(
            (
                _sha256(LOCAL_CUSTOMER_REPORT),
                LOCAL_CUSTOMER_REPORT.stat().st_size,
                LOCAL_CUSTOMER_REPORT.stat().st_mtime_ns,
            ),
            before,
        )

    def _assert_failed_job(self, job) -> None:
        self.assertEqual(job.status, "failed")
        self.assertFalse(job.confirmable)
        self.assertTrue(any(issue.severity == "error" for issue in job.issues))
        self.assertTrue(any(issue.blocks_confirmation for issue in job.issues))
        reloaded = get_report_import(job.id)
        self.assertEqual(reloaded.status, "failed")
        self.assertFalse(reloaded.confirmable)

    @staticmethod
    def _job_count() -> int:
        with database.connect() as connection:
            return int(connection.execute("SELECT COUNT(*) FROM report_import_jobs").fetchone()[0])


if __name__ == "__main__":
    unittest.main()
