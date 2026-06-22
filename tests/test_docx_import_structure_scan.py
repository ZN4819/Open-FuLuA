import os
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "backend"))

from app import database  # noqa: E402
from app.config import settings  # noqa: E402
from app.services.docx_generator import generate_project_docx  # noqa: E402
from app.services.docx_importer import DocxImportPackageError, read_docx_package, scan_docx_structure  # noqa: E402


class DocxImportStructureScanTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)
        self.root = Path(self.temp_dir.name)
        self.original_storage_path = settings.storage_path
        object.__setattr__(settings, "storage_path", self.root / "storage")
        os.environ["FULUA_DATABASE_PATH"] = str(self.root / "test.db")
        database.init_db()

    def tearDown(self) -> None:
        os.environ.pop("FULUA_DATABASE_PATH", None)
        object.__setattr__(settings, "storage_path", self.original_storage_path)

    def test_scans_generated_editable_docx_structure(self) -> None:
        project = self._make_project_with_content()
        path = generate_project_docx(project["id"], "editable")

        scan = scan_docx_structure(path)

        self.assertTrue(scan.has_appendix_title)
        self.assertEqual(scan.suggested_project_name, path.stem)
        self.assertEqual(scan.summary["sections"], 8)
        self.assertEqual(scan.summary["tables"], 8)
        self.assertEqual(scan.summary["errors"], 0)
        self.assertEqual([section.code for section in scan.sections], [f"A-{index}" for index in range(1, 9)])
        self.assertEqual(scan.sections[0].title, "物理和环境安全")
        self.assertEqual(scan.sections[0].table_title, "表A-1物理和环境安全测评结果记录")
        self.assertEqual(scan.sections[0].table_type, "technical")
        self.assertEqual(scan.sections[0].row_count, 2)
        self.assertEqual(scan.sections[4].table_type, "management")
        self.assertEqual(scan.sections[4].row_count, 1)
        self.assertEqual(scan.table_candidates[0].section_code, "A-1")
        self.assertEqual(scan.table_candidates[0].column_count, 8)
        self.assertEqual(scan.table_candidates[4].section_code, "A-5")
        self.assertEqual(scan.table_candidates[4].column_count, 5)

    def test_scans_generated_final_docx_structure(self) -> None:
        project = self._make_project_with_content()
        path = generate_project_docx(project["id"], "final")

        scan = scan_docx_structure(path)

        self.assertTrue(scan.has_appendix_title)
        self.assertEqual(scan.summary["tables"], 8)
        self.assertEqual(scan.summary["errors"], 0)
        self.assertTrue(all(candidate.confidence >= 0.9 for candidate in scan.table_candidates))

    def test_reads_package_relationships_and_media_paths(self) -> None:
        project = self._make_project_with_content()
        path = generate_project_docx(project["id"], "editable")

        package = read_docx_package(path)

        self.assertIn("word/media/", ";".join(package.media_paths))
        self.assertTrue(package.relationships)
        self.assertTrue(any(target.startswith("media/") for target in package.relationships.values()))

    def test_rejects_non_docx_and_corrupted_docx(self) -> None:
        text_path = self.root / "not-docx.txt"
        text_path.write_text("not a docx", encoding="utf-8")
        broken_path = self.root / "broken.docx"
        broken_path.write_bytes(b"not a zip package")

        with self.assertRaisesRegex(DocxImportPackageError, "仅支持"):
            read_docx_package(text_path)
        with self.assertRaisesRegex(DocxImportPackageError, "损坏"):
            read_docx_package(broken_path)

    def test_rejects_docx_package_missing_document_xml(self) -> None:
        path = self.root / "missing-document.docx"
        with zipfile.ZipFile(path, "w") as package:
            package.writestr("[Content_Types].xml", "<Types />")

        with self.assertRaisesRegex(DocxImportPackageError, "word/document.xml"):
            read_docx_package(path)

    def test_reports_missing_sections_and_tables_from_partial_appendix_docx(self) -> None:
        path = self.root / "partial.docx"
        document = Document()
        document.add_paragraph("附录A测评结果记录")
        document.add_paragraph("A-1 物理和环境安全")
        document.add_paragraph("表A-1物理和环境安全测评结果记录")
        table = document.add_table(rows=3, cols=8)
        for index, label in enumerate(["测评单元", "测评对象", "结果记录", "D", "A", "K", "测评对象评分", "测评单元得分"]):
            table.cell(0, index).text = label
        table.cell(1, 0).text = "身份鉴别"
        table.cell(1, 1).text = "机房"
        table.cell(1, 2).text = "现场核查。"
        document.save(path)

        scan = scan_docx_structure(path)
        issue_codes = {issue.code for issue in scan.issues}

        self.assertTrue(scan.has_appendix_title)
        self.assertEqual(scan.sections[0].code, "A-1")
        self.assertEqual(scan.sections[0].row_count, 1)
        self.assertIn("IMPORT_MISSING_SECTION", issue_codes)
        self.assertIn("IMPORT_MISSING_TABLE", issue_codes)
        self.assertGreater(scan.summary["warnings"], 0)
        self.assertGreater(scan.summary["errors"], 0)

    def _make_project_with_content(self):
        project = database.create_project("DOCX 导入扫描测试")
        image = self._create_evidence_image(project["id"], "A-1")
        database.replace_section_rows(
            project_id=project["id"],
            code="A-1",
            rows=[
                {
                    "unit": "身份鉴别",
                    "object_name": "服务器",
                    "record_text": f"查看登录策略，见 [[FIG:{image['id']}]]。",
                    "metric_result": {
                        "d": "√",
                        "a": "√",
                        "k": "/",
                        "object_score": "1.0000",
                        "unit_score": "1.0000",
                    },
                    "cross_references": [
                        {
                            "target_image_id": image["id"],
                            "token": f"[[FIG:{image['id']}]]",
                            "display_text": "图A-1-1",
                        }
                    ],
                },
                {
                    "unit": "身份鉴别",
                    "object_name": "备用服务器",
                    "record_text": "查看备用登录策略。",
                    "metric_result": {
                        "d": "√",
                        "a": "√",
                        "k": "/",
                        "object_score": "1.0000",
                        "unit_score": "1.0000",
                    },
                },
            ],
        )
        database.replace_section_rows(
            project_id=project["id"],
            code="A-5",
            rows=[
                {
                    "unit": "管理制度",
                    "object_name": "制度文件",
                    "record_text": "制度文件已建立并发布。",
                    "metric_result": {
                        "compliance": "符合",
                        "unit_score": "1.0000",
                    },
                }
            ],
        )
        return project

    def _create_evidence_image(self, project_id: int, section_code: str):
        relative_path = Path("uploads") / str(project_id) / section_code / "scan.png"
        absolute_path = settings.storage_path / relative_path
        absolute_path.parent.mkdir(parents=True, exist_ok=True)
        Image.new("RGB", (600, 300), color=(255, 255, 255)).save(absolute_path, dpi=(150, 150))
        return database.create_evidence_image(
            project_id,
            section_code,
            {
                "file_path": relative_path.as_posix(),
                "original_name": "scan.png",
                "caption": "登录策略截图",
                "alt_text": "登录策略截图",
                "pixel_width": 600,
                "pixel_height": 300,
                "dpi_x": 150,
                "dpi_y": 150,
                "display_width_in": 4,
                "display_height_in": 2,
            },
        )


if __name__ == "__main__":
    unittest.main()
