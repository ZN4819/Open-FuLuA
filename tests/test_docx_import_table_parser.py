import os
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "backend"))

from app import database  # noqa: E402
from app.config import settings  # noqa: E402
from app.services.docx_generator import generate_project_docx  # noqa: E402
from app.services.docx_importer import parse_docx_core_tables  # noqa: E402


class DocxImportTableParserTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)
        self.root = Path(self.temp_dir.name)
        self.original_storage_path = settings.storage_path
        object.__setattr__(settings, "storage_path", self.root / "storage")
        os.environ["FULUA_DATABASE_PATH"] = str(self.root / "test.db")
        database.init_db()

    def tearDown(self) -> None:
        os.environ.pop("FULUA_DATABASE_PATH", None)
        object.__setattr__(settings, "storage_path", self.original_storage_path)

    def test_parses_generated_editable_docx_core_rows_and_content_controls(self) -> None:
        project = self._make_project_with_content()
        path = generate_project_docx(project["id"], "editable")

        parsed = parse_docx_core_tables(path)
        a1 = self._section(parsed, "A-1")
        a2 = self._section(parsed, "A-2")
        a5 = self._section(parsed, "A-5")

        self.assertEqual(parsed.summary["assessment_rows"], 3)
        self.assertEqual(parsed.summary["errors"], 0)
        self.assertEqual(parsed.summary["warnings"], 0)
        self.assertEqual(a1.row_count, 2)
        self.assertEqual(a1.rows[0].unit, "身份鉴别")
        self.assertEqual(a1.rows[0].object_name, "服务器")
        self.assertIn("查看登录策略", a1.rows[0].record_text)
        self.assertEqual(a1.rows[0].metric_result.d, "√")
        self.assertEqual(a1.rows[0].metric_result.a, "√")
        self.assertEqual(a1.rows[0].metric_result.k, "/")
        self.assertEqual(a1.rows[0].metric_result.object_score, "0.5000")
        self.assertEqual(a1.rows[0].metric_result.unit_score, "0.5000")
        self.assertEqual(a1.rows[1].unit, "身份鉴别")
        self.assertEqual(a1.rows[1].object_name, "备用服务器")
        self.assertEqual(a2.row_count, 0)
        self.assertEqual(a5.row_count, 1)
        self.assertEqual(a5.rows[0].metric_result.compliance, "符合")
        self.assertEqual(a5.rows[0].metric_result.unit_score, "1.0000")

    def test_parses_generated_final_docx_core_rows(self) -> None:
        project = self._make_project_with_content()
        path = generate_project_docx(project["id"], "final")

        parsed = parse_docx_core_tables(path)
        a1 = self._section(parsed, "A-1")
        a5 = self._section(parsed, "A-5")

        self.assertEqual(parsed.summary["assessment_rows"], 3)
        self.assertEqual(a1.rows[0].metric_result.d, "√")
        self.assertEqual(a1.rows[1].unit, "身份鉴别")
        self.assertEqual(a5.rows[0].metric_result.compliance, "符合")

    def test_inherits_vertically_merged_unit_cells(self) -> None:
        path = self.root / "merged-unit.docx"
        document = Document()
        document.add_paragraph("附录A测评结果记录")
        document.add_paragraph("A-1 物理和环境安全")
        document.add_paragraph("表A-1物理和环境安全测评结果记录")
        table = document.add_table(rows=4, cols=8)
        self._fill_technical_header(table)
        merged = table.cell(2, 0).merge(table.cell(3, 0))
        merged.text = "身份鉴别"
        table.cell(2, 1).text = "主机"
        table.cell(2, 2).text = "主机登录核查。"
        table.cell(2, 3).text = "√"
        table.cell(2, 4).text = "√"
        table.cell(2, 5).text = "/"
        table.cell(2, 6).text = "1.0000"
        table.cell(2, 7).text = "1.0000"
        table.cell(3, 1).text = "备机"
        table.cell(3, 2).text = "备机登录核查。"
        table.cell(3, 3).text = "√"
        table.cell(3, 4).text = "√"
        table.cell(3, 5).text = "/"
        table.cell(3, 6).text = "1.0000"
        table.cell(3, 7).text = "1.0000"
        document.save(path)

        parsed = parse_docx_core_tables(path)
        a1 = self._section(parsed, "A-1")

        self.assertEqual(a1.row_count, 2)
        self.assertEqual([row.unit for row in a1.rows], ["身份鉴别", "身份鉴别"])
        self.assertEqual([row.object_name for row in a1.rows], ["主机", "备机"])

    def test_filters_empty_placeholder_rows(self) -> None:
        path = self.root / "empty-row.docx"
        document = Document()
        document.add_paragraph("附录A测评结果记录")
        document.add_paragraph("A-1 物理和环境安全")
        document.add_paragraph("表A-1物理和环境安全测评结果记录")
        table = document.add_table(rows=4, cols=8)
        self._fill_technical_header(table)
        for column in range(3, 6):
            table.cell(2, column).text = "/"
        table.cell(3, 0).text = "身份鉴别"
        table.cell(3, 1).text = "主机"
        table.cell(3, 2).text = "有效数据行。"
        table.cell(3, 3).text = "√"
        table.cell(3, 4).text = "√"
        table.cell(3, 5).text = "/"
        table.cell(3, 6).text = "1.0000"
        table.cell(3, 7).text = "1.0000"
        document.save(path)

        parsed = parse_docx_core_tables(path)
        a1 = self._section(parsed, "A-1")

        self.assertEqual(a1.row_count, 1)
        self.assertEqual(a1.rows[0].object_name, "主机")

    def test_parses_sdt_wrapped_technical_metric_cells(self) -> None:
        path = self.root / "sdt-technical.docx"
        document = Document()
        document.add_paragraph("A-1")
        table = document.add_table(rows=4, cols=8)
        self._fill_technical_header(table)
        table.cell(2, 0).text = "Unit"
        table.cell(2, 1).text = "Object 1"
        table.cell(2, 2).text = "Record 1"
        table.cell(2, 3).text = "\u221a"
        table.cell(2, 4).text = "\u221a"
        table.cell(2, 5).text = "\u221a"
        table.cell(2, 6).text = "1.0000"
        table.cell(2, 7).text = "0.5000"
        table.cell(3, 0).text = "Unit"
        table.cell(3, 1).text = "Object 2"
        table.cell(3, 2).text = "Record 2"
        table.cell(3, 3).text = "\u00d7"
        table.cell(3, 4).text = "/"
        table.cell(3, 5).text = "/"
        table.cell(3, 6).text = "0.0000"
        table.cell(3, 7).text = "0.5000"
        self._wrap_row_cells_in_sdt(table.rows[2], [3, 4, 5])
        self._wrap_row_cells_in_sdt(table.rows[3], [3, 4, 5])
        document.save(path)

        parsed = parse_docx_core_tables(path)
        a1 = self._section(parsed, "A-1")
        issue_codes = {issue.code for issue in parsed.issues}

        self.assertEqual(a1.row_count, 2)
        self.assertEqual(a1.rows[0].metric_result.d, "\u221a")
        self.assertEqual(a1.rows[0].metric_result.a, "\u221a")
        self.assertEqual(a1.rows[0].metric_result.k, "\u221a")
        self.assertEqual(a1.rows[0].metric_result.object_score, "1.0000")
        self.assertEqual(a1.rows[0].metric_result.unit_score, "0.5000")
        self.assertEqual(a1.rows[1].metric_result.d, "\u00d7")
        self.assertEqual(a1.rows[1].metric_result.a, "/")
        self.assertEqual(a1.rows[1].metric_result.k, "/")
        self.assertEqual(a1.rows[1].metric_result.object_score, "0.0000")
        self.assertEqual(a1.rows[1].metric_result.unit_score, "0.5000")
        self.assertNotIn("IMPORT_INVALID_DAK_VALUE", issue_codes)

    def test_reports_invalid_metric_and_compliance_values(self) -> None:
        technical_path = self.root / "invalid-technical.docx"
        document = Document()
        document.add_paragraph("附录A测评结果记录")
        document.add_paragraph("A-1 物理和环境安全")
        document.add_paragraph("表A-1物理和环境安全测评结果记录")
        table = document.add_table(rows=3, cols=8)
        self._fill_technical_header(table)
        table.cell(2, 0).text = "身份鉴别"
        table.cell(2, 1).text = "主机"
        table.cell(2, 2).text = "指标值异常。"
        table.cell(2, 3).text = "是"
        table.cell(2, 4).text = "√"
        table.cell(2, 5).text = "/"
        document.save(technical_path)

        management_path = self.root / "invalid-management.docx"
        document = Document()
        document.add_paragraph("附录A测评结果记录")
        document.add_paragraph("A-5 管理制度")
        document.add_paragraph("表A-5管理制度测评结果记录")
        table = document.add_table(rows=2, cols=5)
        for index, label in enumerate(["测评单元", "测评对象", "结果记录", "符合情况", "测评单元得分"]):
            table.cell(0, index).text = label
        table.cell(1, 0).text = "管理制度"
        table.cell(1, 1).text = "制度文件"
        table.cell(1, 2).text = "符合情况异常。"
        table.cell(1, 3).text = "未知"
        document.save(management_path)

        technical_codes = {issue.code for issue in parse_docx_core_tables(technical_path).issues}
        management_codes = {issue.code for issue in parse_docx_core_tables(management_path).issues}

        self.assertIn("IMPORT_INVALID_DAK_VALUE", technical_codes)
        self.assertIn("IMPORT_INVALID_COMPLIANCE_VALUE", management_codes)

    def _make_project_with_content(self):
        project = database.create_project("DOCX 表格解析测试")
        image = self._create_evidence_image(project["id"], "A-1")
        database.replace_section_rows(
            project_id=project["id"],
            code="A-1",
            rows=[
                {
                    "unit": "身份鉴别",
                    "object_name": "服务器",
                    "record_text": f"查看登录策略，见 [[FIG:{image['id']}]]。",
                    "metric_result": {
                        "d": "√",
                        "a": "√",
                        "k": "/",
                        "object_score": "1.0000",
                        "unit_score": "1.0000",
                    },
                    "cross_references": [
                        {
                            "target_image_id": image["id"],
                            "token": f"[[FIG:{image['id']}]]",
                            "display_text": "图A-1-1",
                        }
                    ],
                },
                {
                    "unit": "身份鉴别",
                    "object_name": "备用服务器",
                    "record_text": "查看备用登录策略。",
                    "metric_result": {
                        "d": "√",
                        "a": "√",
                        "k": "/",
                        "object_score": "1.0000",
                        "unit_score": "1.0000",
                    },
                },
            ],
        )
        database.replace_section_rows(
            project_id=project["id"],
            code="A-5",
            rows=[
                {
                    "unit": "管理制度",
                    "object_name": "制度文件",
                    "record_text": "制度文件已建立并发布。",
                    "metric_result": {
                        "compliance": "符合",
                        "unit_score": "1.0000",
                    },
                }
            ],
        )
        return project

    def _create_evidence_image(self, project_id: int, section_code: str):
        relative_path = Path("uploads") / str(project_id) / section_code / "parser.png"
        absolute_path = settings.storage_path / relative_path
        absolute_path.parent.mkdir(parents=True, exist_ok=True)
        Image.new("RGB", (600, 300), color=(255, 255, 255)).save(absolute_path, dpi=(150, 150))
        return database.create_evidence_image(
            project_id,
            section_code,
            {
                "file_path": relative_path.as_posix(),
                "original_name": "parser.png",
                "caption": "登录策略截图",
                "alt_text": "登录策略截图",
                "pixel_width": 600,
                "pixel_height": 300,
                "dpi_x": 150,
                "dpi_y": 150,
                "display_width_in": 4,
                "display_height_in": 2,
            },
        )

    @staticmethod
    def _wrap_row_cells_in_sdt(row, indexes: list[int]) -> None:
        tr = row._tr
        for index in sorted(indexes, reverse=True):
            cell = tr.tc_lst[index]
            position = list(tr).index(cell)
            tr.remove(cell)
            sdt = OxmlElement("w:sdt")
            content = OxmlElement("w:sdtContent")
            content.append(cell)
            sdt.append(content)
            tr.insert(position, sdt)

    @staticmethod
    def _fill_technical_header(table) -> None:
        for index, label in enumerate(["测评单元", "测评对象", "结果记录", "D", "A", "K", "测评对象评分", "测评单元得分"]):
            table.cell(0, index).text = label
        for index, label in enumerate(["", "", "", "D", "A", "K", "测评对象评分", ""]):
            table.cell(1, index).text = label

    @staticmethod
    def _section(parsed, code: str):
        return next(section for section in parsed.sections if section.code == code)


if __name__ == "__main__":
    unittest.main()
