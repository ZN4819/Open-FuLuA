import json
import os
import sys
import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import HTTPException, UploadFile
from PIL import Image
from starlette.datastructures import Headers


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "backend"))

from app import database  # noqa: E402
from app.api.imports import create_project_from_docx_import as api_confirm_docx_import  # noqa: E402
from app.api.imports import upload_docx_import as api_upload_docx_import  # noqa: E402
from app.api.sections import build_section_detail  # noqa: E402
from app.config import settings  # noqa: E402
from app.schemas import DocxImportCreateProjectRequest  # noqa: E402
from app.services.docx_generator import generate_project_docx  # noqa: E402
from app.services.validator import validate_project  # noqa: E402


class DocxImportConfirmApiTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)
        self.root = Path(self.temp_dir.name)
        self.original_storage_path = settings.storage_path
        object.__setattr__(settings, "storage_path", self.root / "storage")
        os.environ["FULUA_DATABASE_PATH"] = str(self.root / "test.db")
        database.init_db()

    def tearDown(self) -> None:
        os.environ.pop("FULUA_DATABASE_PATH", None)
        object.__setattr__(settings, "storage_path", self.original_storage_path)

    def test_confirm_import_creates_project_images_rows_references_and_validates(self) -> None:
        source_project = self._make_project_with_content()
        exported = generate_project_docx(source_project["id"], "editable")
        preview = api_upload_docx_import(self._upload("原系统附录A.docx", exported.read_bytes()))

        confirmed = api_confirm_docx_import(
            preview.id,
            DocxImportCreateProjectRequest(project_name="导入后的系统附录A"),
        )

        self.assertEqual(confirmed.status, "succeeded")
        self.assertIsNotNone(confirmed.created_project_id)
        imported_project = database.get_project_by_id(int(confirmed.created_project_id))
        self.assertEqual(imported_project["name"], "导入后的系统附录A")

        detail = build_section_detail(int(confirmed.created_project_id), "A-1")
        self.assertEqual(len(detail.rows), 1)
        self.assertEqual(len(detail.evidence_images), 1)
        self.assertEqual(len(detail.cross_references), 1)
        image = detail.evidence_images[0]
        row = detail.rows[0]
        reference = detail.cross_references[0]

        self.assertTrue((settings.storage_path / image.file_path).exists())
        self.assertTrue(image.file_path.startswith(f"uploads/{confirmed.created_project_id}/A-1/"))
        self.assertEqual(row.unit, "身份鉴别")
        self.assertEqual(row.metric_result.d, "√")
        self.assertIn(f"[[FIG:{image.id}]]", row.record_text)
        self.assertNotIn("[[FIG:import:", row.record_text)
        self.assertEqual(reference.target_image_id, image.id)
        self.assertEqual(reference.token, f"[[FIG:{image.id}]]")
        self.assertEqual(image.project_image_no, 1)
        self.assertEqual(image.figure_label, "图A-1-1")
        self.assertEqual(reference.display_text, "图A-1-1")

        validation = validate_project(int(confirmed.created_project_id))
        self.assertEqual(validation.summary.errors, 0)

    def test_confirm_import_rejects_preview_with_errors_without_creating_project(self) -> None:
        preview = api_upload_docx_import(self._upload("partial.docx", self._partial_docx_bytes()))
        self.assertFalse(preview.can_create_project)
        before_projects = len(database.list_projects())

        with self.assertRaises(HTTPException) as context:
            api_confirm_docx_import(preview.id, DocxImportCreateProjectRequest(project_name="不应创建"))

        self.assertEqual(context.exception.status_code, 400)
        self.assertIn("包含错误", str(context.exception.detail))
        self.assertEqual(len(database.list_projects()), before_projects)
        self.assertEqual(database.get_docx_import_job(preview.id)["status"], "preview_ready")

    def test_confirm_import_rolls_back_project_when_import_image_is_missing(self) -> None:
        source_project = self._make_project_with_content()
        exported = generate_project_docx(source_project["id"], "editable")
        preview = api_upload_docx_import(self._upload("缺失图片附录A.docx", exported.read_bytes()))
        parsed_path = settings.storage_path / str(preview.parsed_json_path)
        parsed_payload = json.loads(parsed_path.read_text(encoding="utf-8"))
        import_image_path = settings.storage_path / parsed_payload["sections"][0]["images"][0]["file_path"]
        import_image_path.unlink()
        before_projects = len(database.list_projects())
        next_project_id = max(project["id"] for project in database.list_projects()) + 1

        with self.assertRaises(HTTPException) as context:
            api_confirm_docx_import(preview.id, DocxImportCreateProjectRequest(project_name="缺失图片导入"))

        self.assertEqual(context.exception.status_code, 400)
        self.assertIn("导入图片文件不存在", str(context.exception.detail))
        self.assertEqual(len(database.list_projects()), before_projects)
        self.assertFalse((settings.storage_path / "uploads" / str(next_project_id)).exists())
        failed_job = database.get_docx_import_job(preview.id)
        self.assertEqual(failed_job["status"], "failed")
        self.assertIsNone(failed_job["created_project_id"])

    def _make_project_with_content(self):
        project = database.create_project("DOCX 确认导入测试源项目")
        image = self._create_evidence_image(project["id"], "A-1")
        database.replace_section_rows(
            project_id=project["id"],
            code="A-1",
            rows=[
                {
                    "unit": "身份鉴别",
                    "object_name": "服务器",
                    "record_text": f"查看登录策略，见 [[FIG:{image['id']}]]。",
                    "metric_result": {
                        "d": "√",
                        "a": "√",
                        "k": "/",
                        "object_score": "1.0000",
                        "unit_score": "1.0000",
                    },
                    "cross_references": [
                        {
                            "target_image_id": image["id"],
                            "token": f"[[FIG:{image['id']}]]",
                            "display_text": "图A-1-1",
                        }
                    ],
                }
            ],
        )
        return project

    def _create_evidence_image(self, project_id: int, section_code: str):
        relative_path = Path("uploads") / str(project_id) / section_code / "preview.png"
        absolute_path = settings.storage_path / relative_path
        absolute_path.parent.mkdir(parents=True, exist_ok=True)
        Image.new("RGB", (600, 300), color=(255, 255, 255)).save(absolute_path, dpi=(150, 150))
        return database.create_evidence_image(
            project_id,
            section_code,
            {
                "file_path": relative_path.as_posix(),
                "original_name": "preview.png",
                "caption": "登录策略截图",
                "alt_text": "登录策略截图",
                "pixel_width": 600,
                "pixel_height": 300,
                "dpi_x": 150,
                "dpi_y": 150,
                "display_width_in": 4,
                "display_height_in": 2,
            },
        )

    def _partial_docx_bytes(self) -> bytes:
        path = self.root / "partial.docx"
        document = Document()
        document.add_paragraph("附录A测评结果记录")
        document.add_paragraph("A-1 物理和环境安全")
        document.add_paragraph("表A-1物理和环境安全测评结果记录")
        table = document.add_table(rows=2, cols=8)
        for index, label in enumerate(["测评单元", "测评对象", "结果记录", "D", "A", "K", "测评对象评分", "测评单元得分"]):
            table.cell(0, index).text = label
        table.cell(1, 0).text = "身份鉴别"
        table.cell(1, 1).text = "机房"
        table.cell(1, 2).text = "现场查看。"
        table.cell(1, 3).text = "√"
        table.cell(1, 4).text = "√"
        table.cell(1, 5).text = "/"
        path.parent.mkdir(parents=True, exist_ok=True)
        document.save(path)
        return path.read_bytes()

    @staticmethod
    def _upload(filename: str, content: bytes, content_type: str = "application/vnd.openxmlformats-officedocument.wordprocessingml.document") -> UploadFile:
        return UploadFile(
            file=BytesIO(content),
            filename=filename,
            headers=Headers({"content-type": content_type}),
        )


if __name__ == "__main__":
    unittest.main()
