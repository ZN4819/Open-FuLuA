from __future__ import annotations

import json
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "backend"))

from app.services.report_templates.analyzer import UnsafePackageError, analyze_report_template


class ReportTemplateAnalyzerTests(unittest.TestCase):
    BASE_TEMPLATE = ROOT / "02-商用密码应用安全性评估报告模板(2023版)—系统密评报告--三级-20251208修改.docx"
    BASELINE = ROOT / "tests" / "fixtures" / "report_templates" / "base_template_baseline.json"

    def _fixture(self, directory: Path) -> Path:
        path = directory / "fixture.docx"
        document = Document()
        document.add_paragraph("合成测试，不含客户数据")
        document.add_table(rows=2, cols=2)
        document.add_section()
        document.save(path)
        return path

    def test_analyzes_synthetic_docx_without_disclosing_path_or_text(self) -> None:
        with tempfile.TemporaryDirectory() as value:
            path = self._fixture(Path(value))
            result = analyze_report_template(path, source_role="synthetic_fixture")
        payload = json.dumps(result.model_dump(mode="json"), ensure_ascii=False)
        self.assertEqual(result.document.table_count, 1)
        self.assertEqual(result.document.section_count, 2)
        self.assertNotIn("fixture.docx", payload)
        self.assertNotIn("合成测试", payload)

    def test_rejects_path_traversal_member(self) -> None:
        with tempfile.TemporaryDirectory() as value:
            path = Path(value) / "unsafe.docx"
            with zipfile.ZipFile(path, "w") as package:
                package.writestr("../word/document.xml", b"x")
            with self.assertRaisesRegex(UnsafePackageError, "ZIP_MEMBER_PATH_INVALID"):
                analyze_report_template(path, source_role="synthetic_fixture")

    def test_rejects_case_insensitive_duplicate_members(self) -> None:
        with tempfile.TemporaryDirectory() as value:
            path = Path(value) / "unsafe.docx"
            with zipfile.ZipFile(path, "w") as package:
                package.writestr("word/document.xml", b"x")
                package.writestr("WORD/document.xml", b"x")
            with self.assertRaisesRegex(UnsafePackageError, "ZIP_MEMBER_DUPLICATE"):
                analyze_report_template(path, source_role="synthetic_fixture")

    def test_rejects_dtd_before_xml_parse(self) -> None:
        with tempfile.TemporaryDirectory() as value:
            path = self._fixture(Path(value))
            rewritten = Path(value) / "unsafe.docx"
            with zipfile.ZipFile(path) as source, zipfile.ZipFile(rewritten, "w") as target:
                for info in source.infolist():
                    data = source.read(info)
                    if info.filename == "word/document.xml":
                        data = b'<!DOCTYPE x [<!ENTITY y "z">]><x>&y;</x>'
                    target.writestr(info, data)
            with self.assertRaisesRegex(UnsafePackageError, "XML_DTD_OR_ENTITY_FORBIDDEN"):
                analyze_report_template(rewritten, source_role="synthetic_fixture")

    @unittest.skipUnless(BASE_TEMPLATE.exists(), "只在本地只读源模板存在时运行")
    def test_approved_base_template_matches_committed_forensics_baseline(self) -> None:
        before = (self.BASE_TEMPLATE.stat().st_size, self.BASE_TEMPLATE.stat().st_mtime_ns)
        actual = analyze_report_template(self.BASE_TEMPLATE, source_role="base_template")
        expected = json.loads(self.BASELINE.read_text(encoding="utf-8"))
        self.assertEqual(actual.model_dump(mode="json"), expected)
        self.assertEqual(
            actual.source_sha256,
            "b3957fd1da3bf19c31ac515fbdc6bf989fd7df033ca4d179c4b6e9567247fcf8",
        )
        self.assertEqual((self.BASE_TEMPLATE.stat().st_size, self.BASE_TEMPLATE.stat().st_mtime_ns), before)


if __name__ == "__main__":
    unittest.main()
