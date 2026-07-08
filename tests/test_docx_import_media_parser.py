import os
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "backend"))

from app import database  # noqa: E402
from app.config import settings  # noqa: E402
from app.services.docx_generator import generate_project_docx  # noqa: E402
from app.services.docx_importer import parse_docx_images_and_references  # noqa: E402


class DocxImportMediaParserTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)
        self.root = Path(self.temp_dir.name)
        self.original_storage_path = settings.storage_path
        object.__setattr__(settings, "storage_path", self.root / "storage")
        os.environ["FULUA_DATABASE_PATH"] = str(self.root / "test.db")
        database.init_db()

    def tearDown(self) -> None:
        os.environ.pop("FULUA_DATABASE_PATH", None)
        object.__setattr__(settings, "storage_path", self.original_storage_path)

    def test_extracts_generated_images_captions_order_and_ref_tokens(self) -> None:
        project = self._make_project_with_two_images()
        path = generate_project_docx(project["id"], "editable")
        import_dir = settings.storage_path / "imports" / "1"

        parsed = parse_docx_images_and_references(path, import_dir)
        a1 = self._section(parsed, "A-1")

        self.assertEqual(parsed.summary["images"], 2)
        self.assertEqual(parsed.summary["references"], 1)
        self.assertEqual(a1.image_count, 2)
        self.assertEqual([image.figure_label for image in a1.images], ["图A-1-1", "图A-1-2"])
        self.assertEqual([image.caption for image in a1.images], ["登录策略截图", "备用策略截图"])
        self.assertEqual([image.sort_order for image in a1.images], [1, 2])
        self.assertTrue((settings.storage_path / a1.images[0].file_path).exists())
        self.assertEqual(a1.images[0].pixel_width, 600)
        self.assertEqual(a1.rows[0].record_text, "查看登录策略，见 [[FIG:import:A-1-1]]。")
        self.assertEqual(a1.rows[0].cross_references[0].display_text, "图A-1-1")
        self.assertEqual(a1.rows[0].cross_references[0].target_image_key, "A-1-1")

    def test_maps_plain_visible_figure_numbers_to_import_tokens(self) -> None:
        path = self.root / "visible-reference.docx"
        image_path = self._local_image("visible.png", (320, 160))
        document = Document()
        self._add_technical_section_with_row(document, "现场查看，见图A-1-1。")
        document.add_picture(str(image_path))
        document.add_paragraph("图A-1-1 普通可见图号截图")
        document.save(path)

        parsed = parse_docx_images_and_references(path, settings.storage_path / "imports" / "2")
        a1 = self._section(parsed, "A-1")

        self.assertEqual(a1.image_count, 1)
        self.assertEqual(a1.images[0].caption, "普通可见图号截图")
        self.assertEqual(a1.rows[0].record_text, "现场查看，见[[FIG:import:A-1-1]]。")
        self.assertEqual(a1.rows[0].cross_references[0].token, "[[FIG:import:A-1-1]]")

    def test_combines_multiple_images_that_share_one_caption(self) -> None:
        path = self.root / "shared-caption.docx"
        first_image = self._local_image("shared-first.png", (120, 80), color=(255, 0, 0))
        second_image = self._local_image("shared-second.png", (80, 100), color=(0, 0, 255))
        document = Document()
        self._add_technical_section_with_row(document, "现场查看，见图A-1-1。")
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(str(first_image))
        run.add_picture(str(second_image))
        document.add_paragraph("图A-1-1 组合证据图")
        document.save(path)

        parsed = parse_docx_images_and_references(path, settings.storage_path / "imports" / "shared")
        a1 = self._section(parsed, "A-1")
        issue_codes = {issue.code for issue in parsed.issues}

        self.assertEqual(a1.image_count, 1)
        self.assertEqual(parsed.summary["images"], 1)
        self.assertEqual(a1.images[0].figure_label, "图A-1-1")
        self.assertEqual(a1.images[0].caption, "组合证据图")
        self.assertEqual(len(a1.images[0].relationship_id.split(",")), 2)
        self.assertIn("IMPORT_IMAGE_CAPTION_SHARED", issue_codes)
        self.assertNotIn("IMPORT_IMAGE_CAPTION_DUPLICATE", issue_codes)
        combined_path = settings.storage_path / a1.images[0].file_path
        self.assertTrue(combined_path.exists())
        with Image.open(combined_path) as combined:
            self.assertEqual(combined.size, (200, 100))
        self.assertEqual(a1.rows[0].record_text, "现场查看，见[[FIG:import:A-1-1]]。")
        self.assertEqual(a1.rows[0].cross_references[0].target_image_key, "A-1-1")

    def test_reports_missing_caption_without_dropping_image(self) -> None:
        path = self.root / "missing-caption.docx"
        image_path = self._local_image("no-caption.png", (300, 200))
        document = Document()
        self._add_technical_section_with_row(document, "现场查看设备。")
        document.add_picture(str(image_path))
        document.save(path)

        parsed = parse_docx_images_and_references(path, settings.storage_path / "imports" / "3")
        a1 = self._section(parsed, "A-1")
        issue_codes = {issue.code for issue in parsed.issues}

        self.assertEqual(a1.image_count, 1)
        self.assertIn("IMPORT_IMAGE_CAPTION_MISSING", issue_codes)
        self.assertTrue((settings.storage_path / a1.images[0].file_path).exists())

    def test_reports_broken_visible_figure_reference(self) -> None:
        path = self.root / "broken-reference.docx"
        image_path = self._local_image("broken.png", (300, 200))
        document = Document()
        self._add_technical_section_with_row(document, "现场查看，见图A-1-99。")
        document.add_picture(str(image_path))
        document.add_paragraph("图A-1-1 实际图片")
        document.save(path)

        parsed = parse_docx_images_and_references(path, settings.storage_path / "imports" / "4")
        a1 = self._section(parsed, "A-1")
        issue_codes = {issue.code for issue in parsed.issues}

        self.assertEqual(a1.image_count, 1)
        self.assertEqual(a1.rows[0].record_text, "现场查看，见图A-1-99。")
        self.assertEqual(a1.rows[0].cross_references, [])
        self.assertIn("IMPORT_REFERENCE_TARGET_MISSING", issue_codes)

    def _make_project_with_two_images(self):
        project = database.create_project("DOCX 媒体解析测试")
        first_image = self._create_evidence_image(project["id"], "A-1", "first.png", "登录策略截图", (600, 300))
        second_image = self._create_evidence_image(project["id"], "A-1", "second.png", "备用策略截图", (400, 300))
        database.replace_section_rows(
            project_id=project["id"],
            code="A-1",
            rows=[
                {
                    "unit": "身份鉴别",
                    "object_name": "服务器",
                    "record_text": f"查看登录策略，见 [[FIG:{first_image['id']}]]。",
                    "metric_result": {
                        "d": "√",
                        "a": "√",
                        "k": "/",
                        "object_score": "1.0000",
                        "unit_score": "1.0000",
                    },
                    "cross_references": [
                        {
                            "target_image_id": first_image["id"],
                            "token": f"[[FIG:{first_image['id']}]]",
                            "display_text": "图A-1-1",
                        }
                    ],
                }
            ],
        )
        second_image  # 保持第二张图不被正文引用，用于验证图片排序恢复。
        return project

    def _create_evidence_image(self, project_id: int, section_code: str, filename: str, caption: str, size: tuple[int, int]):
        relative_path = Path("uploads") / str(project_id) / section_code / filename
        absolute_path = settings.storage_path / relative_path
        absolute_path.parent.mkdir(parents=True, exist_ok=True)
        Image.new("RGB", size, color=(255, 255, 255)).save(absolute_path, dpi=(150, 150))
        return database.create_evidence_image(
            project_id,
            section_code,
            {
                "file_path": relative_path.as_posix(),
                "original_name": filename,
                "caption": caption,
                "alt_text": caption,
                "pixel_width": size[0],
                "pixel_height": size[1],
                "dpi_x": 150,
                "dpi_y": 150,
                "display_width_in": 4,
                "display_height_in": 2,
            },
        )

    def _local_image(self, filename: str, size: tuple[int, int], color: tuple[int, int, int] = (255, 255, 255)) -> Path:
        path = self.root / filename
        Image.new("RGB", size, color=color).save(path, dpi=(144, 144))
        return path

    @staticmethod
    def _add_technical_section_with_row(document: Document, record_text: str) -> None:
        document.add_paragraph("附录A测评结果记录")
        document.add_paragraph("A-1 物理和环境安全")
        document.add_paragraph("表A-1物理和环境安全测评结果记录")
        table = document.add_table(rows=2, cols=8)
        for index, label in enumerate(["测评单元", "测评对象", "结果记录", "D", "A", "K", "测评对象评分", "测评单元得分"]):
            table.cell(0, index).text = label
        table.cell(1, 0).text = "身份鉴别"
        table.cell(1, 1).text = "机房"
        table.cell(1, 2).text = record_text
        table.cell(1, 3).text = "√"
        table.cell(1, 4).text = "√"
        table.cell(1, 5).text = "/"
        table.cell(1, 6).text = "1.0000"
        table.cell(1, 7).text = "1.0000"

    @staticmethod
    def _section(parsed, code: str):
        return next(section for section in parsed.sections if section.code == code)


if __name__ == "__main__":
    unittest.main()
